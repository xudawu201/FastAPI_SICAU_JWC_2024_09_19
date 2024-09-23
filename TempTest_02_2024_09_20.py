'''
Author: xudawu
Date: 2024-09-20 17:40:30
LastEditors: xudawu
LastEditTime: 2024-09-20 17:40:35
'''
from docx import Document
from docx.shared import Pt

def create_word_document():
    # 创建一个新的Word文档对象
    doc = Document()

    # 添加主标题，并设置其属性
    main_title = doc.add_paragraph()
    run = main_title.add_run('这是主标题')
    run.font.name = '宋体'
    run.font.size = Pt(18)  # 字体大小二号，近似为18点
    main_title.alignment = 1  # 居中对齐

    # 添加一个段落间隔
    doc.add_paragraph()

    # 添加一个小标题，并设置其属性
    sub_title = doc.add_paragraph()
    sub_run = sub_title.add_run('这是小标题')
    sub_run.font.name = '宋体'
    sub_run.font.size = Pt(14)  # 字体大小四号，近似为14点
    sub_title.style = 'Heading 1'  # 设置为默认的Heading 1样式，可按需调整

    # 添加一个段落间隔
    doc.add_paragraph()

    # 添加正文文本
    body_text = doc.add_paragraph('这是正文内容，这里可以添加更多的信息，比如讨论、数据等。')
    body_text.runs[0].font.name = '宋体'
    body_text.runs[0].font.size = Pt(12)  # 正文一般使用12点字体

    # 保存文档
    doc.save('demo.docx')

create_word_document()