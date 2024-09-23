'''
Author: xudawu
Date: 2024-09-20 09:11:00
LastEditors: xudawu
LastEditTime: 2024-09-20 17:34:02
'''
from docx import Document

def create_word_document():
    # 创建一个新的Word文档对象
    doc = Document()

    # 添加一段文字
    doc.add_paragraph('Hello, this is a test document created with python-docx.')

    # 添加一个标题
    doc.add_heading('Document Title', 0)

    # 保存文档
    doc.save('demo.docx')

create_word_document()